#!/usr/bin/env python3
"""
COI System Document Parser
Extracts text content from PDF, Excel, and DOCX files in the COI System folder.
"""

import os
import sys
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    print("Error: pdfplumber not installed. Run: pip install -r requirements.txt")
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("Error: openpyxl not installed. Run: pip install -r requirements.txt")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install -r requirements.txt")
    sys.exit(1)

try:
    import pandas as pd
except ImportError:
    print("Warning: pandas not installed. Excel parsing will use openpyxl only.")
    pd = None


def parse_pdf(file_path):
    """Extract text from PDF file."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    text_content.append(f"--- Page {page_num} ---\n{text}\n")
        return "\n".join(text_content)
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"


def parse_excel(file_path):
    """Extract data from Excel file."""
    content = []
    try:
        if pd:
            # Use pandas for better data handling
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                content.append(f"\n--- Sheet: {sheet_name} ---\n")
                content.append(df.to_string())
        else:
            # Use openpyxl as fallback
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"\n--- Sheet: {sheet_name} ---\n")
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    content.append("\t".join(row_data))
        return "\n".join(content)
    except Exception as e:
        return f"Error parsing Excel: {str(e)}"


def parse_docx(file_path):
    """Extract text from DOCX file."""
    text_content = []
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            text_content.append("\n--- Table ---\n")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_content.append("\t".join(row_data))
        
        return "\n".join(text_content)
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"


def parse_coi_documents(coi_folder_path=None, output_dir=None):
    """
    Parse all documents in the COI System folder.
    
    Args:
        coi_folder_path: Path to COI System folder (default: workspace/docs/coi-system)
        output_dir: Directory to save extracted text files (default: None, prints to console)
    """
    # Determine workspace root (parent of scripts folder)
    script_dir = Path(__file__).parent
    workspace_root = script_dir.parent
    
    if coi_folder_path is None:
        coi_folder_path = workspace_root / "docs" / "coi-system"
    
    if not os.path.exists(coi_folder_path):
        print(f"Error: COI System folder not found at {coi_folder_path}")
        return
    
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    # Supported file extensions
    pdf_files = []
    excel_files = []
    docx_files = []
    
    # Scan for files
    for file_path in Path(coi_folder_path).iterdir():
        if file_path.is_file():
            ext = file_path.suffix.lower()
            if ext == '.pdf':
                pdf_files.append(file_path)
            elif ext in ['.xlsx', '.xls']:
                excel_files.append(file_path)
            elif ext == '.docx':
                docx_files.append(file_path)
    
    print(f"Found {len(pdf_files)} PDF files, {len(excel_files)} Excel files, {len(docx_files)} DOCX files\n")
    
    # Parse PDF files
    for pdf_file in pdf_files:
        print(f"Parsing PDF: {pdf_file.name}")
        content = parse_pdf(pdf_file)
        
        if output_dir:
            output_file = Path(output_dir) / f"{pdf_file.stem}.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"  -> Saved to {output_file}")
        else:
            print(f"\n{content}\n{'='*80}\n")
    
    # Parse Excel files
    for excel_file in excel_files:
        print(f"Parsing Excel: {excel_file.name}")
        content = parse_excel(excel_file)
        
        if output_dir:
            output_file = Path(output_dir) / f"{excel_file.stem}.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"  -> Saved to {output_file}")
        else:
            print(f"\n{content}\n{'='*80}\n")
    
    # Parse DOCX files
    for docx_file in docx_files:
        print(f"Parsing DOCX: {docx_file.name}")
        content = parse_docx(docx_file)
        
        if output_dir:
            output_file = Path(output_dir) / f"{docx_file.stem}.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"  -> Saved to {output_file}")
        else:
            print(f"\n{content}\n{'='*80}\n")
    
    print("\nParsing complete!")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Parse COI System documents")
    parser.add_argument(
        "--coi-folder",
        type=str,
        default=None,
        help="Path to COI System folder (default: workspace/docs/coi-system)"
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        default=None,
        help="Directory to save extracted text files (default: print to console)"
    )
    
    args = parser.parse_args()
    
    parse_coi_documents(
        coi_folder_path=args.coi_folder,
        output_dir=args.output_dir
    )


